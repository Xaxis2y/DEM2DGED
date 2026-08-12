#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
update_manual_v0.46.py -- bring DEM2DGED_User_Manual.docx up to v0.46.

Version: 0.46

WHY THIS EXISTS
---------------
The user manual had drifted four releases behind the code: its cover read
"VERSION 0.40-beta", its footer read "v0.40-beta", and its version-history
table stopped at v0.40-beta while dem2dged_lib.VERSION was already "0.46".
Nothing in the project checked that -- audit_pure.py enforces version
consistency across the twelve *code* declarations, and neither it nor
RELEASE_CHECK looks at README.md, QUICKSTART.html or this manual. So the
documents could fall arbitrarily far behind without any gate noticing.

This script performs the edit as a repeatable, reviewable operation rather
than a manual pass through Word, so the same update can be re-run against a
fresh copy of the manual and produce the same result.

WHAT IT CHANGES
---------------
  1. Cover block          "VERSION 0.40-beta" -> "VERSION 0.46-beta",
                          "July 2026" -> "August 2026"
  2. Page footer          "v0.40-beta" -> "v0.46-beta"
  3. Package contents     adds the modules and directories added since v0.40
                          (dem2dged_compare.py, dem2dged_env.py, tests/,
                          audit_pure.py, RELEASE_CHECK, START_HERE.md, ...)
  4. Troubleshooting      adds the three failure modes fixed in v0.42-v0.46:
                          wrong Python interpreter, non-UTF-8 console,
                          source CRS with no EPSG code
  5. Version history      adds rows for v0.41 through v0.46

Every step is guarded: re-running the script is a no-op, not a duplicate.
Formatting is preserved by deep-copying an existing paragraph / table row of
the right kind and replacing only its text, so no style needs to be
reconstructed by hand.

USAGE (Anaconda Prompt, dedicated environment -- never base):
    conda activate dem2dged_anaconda_environment
    pip install python-docx
    python update_manual_v0.46.py

Note the command form: "python update_manual_v0.46.py", not
"update_manual_v0.46.py". See dem2dged_env.py for why that matters.
"""

import copy
import os
import shutil
import sys
from datetime import datetime

try:
    from docx import Document
except ImportError:
    sys.stderr.write(
        "ERROR: python-docx is not available in the interpreter that is running\n"
        "       this script (%s).\n"
        "       Install it with:  pip install python-docx\n"
        "       If you believe it IS installed, run 'python dem2dged_env.py'\n"
        "       first -- you are probably on a different interpreter than you\n"
        "       think.\n" % sys.executable
    )
    raise SystemExit(2)

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(HERE, "DEM2DGED_User_Manual.docx")
NEW_VERSION = "0.46-beta"
NEW_DATE = "August 2026"

# Table indices, confirmed by inspection of the v0.40 manual.
T_COVER_VERSION = 2
T_PACKAGE_CONTENTS = 5
T_VERSION_HISTORY = 30


# -- helpers ------------------------------------------------------------------

def set_cell_text(cell, text):
    """Replace a cell's text, keeping the formatting of its first run."""
    para = cell.paragraphs[0]
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for run in para.runs[1:]:
        run._r.getparent().remove(run._r)
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)


def clone_row_after(table, template_index, insert_after_index, values):
    """Deep-copy a row for its formatting, fill it, and place it in the table."""
    new_tr = copy.deepcopy(table.rows[template_index]._tr)
    table.rows[insert_after_index]._tr.addnext(new_tr)
    row = table.rows[insert_after_index + 1]
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value)
    return row


def append_row(table, template_index, values):
    return clone_row_after(table, template_index, len(table.rows) - 1, values)


def build_paragraph(template_p, segments, body_run_idx, code_run_idx,
                    lead_run_idx=None, lead_text=None):
    """Clone a paragraph and rebuild its runs from (text, kind) segments.

    kind is "body" for prose or "code" for a monospaced inline snippet; the
    run formatting for each is taken from the template paragraph itself, so
    the result matches the surrounding document exactly.
    """
    new_p = copy.deepcopy(template_p._p)
    runs = list(new_p.findall(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"))
    body_tpl = copy.deepcopy(runs[body_run_idx])
    code_tpl = copy.deepcopy(runs[code_run_idx])
    lead_tpl = copy.deepcopy(runs[lead_run_idx]) if lead_run_idx is not None else None
    for run in runs:
        new_p.remove(run)

    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    def emit(template, text):
        run = copy.deepcopy(template)
        t = run.find(W + "t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        new_p.append(run)

    if lead_tpl is not None and lead_text is not None:
        emit(lead_tpl, lead_text)
    for text, kind in segments:
        emit(code_tpl if kind == "code" else body_tpl, text)
    return new_p


def para_index_by_text(doc, needle):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == needle:
            return i
    return -1


# -- content ------------------------------------------------------------------

PACKAGE_ROWS = [
    ("dem2dged_compare.py",
     "Resampling comparison test — ranks Nearest / Bilinear / Cubic against the "
     "source and writes an HTML report"),
    ("dem2dged_env.py",
     "Environment diagnostic (v0.46) — dependency-free; run it when an import "
     "fails to see which interpreter is actually running"),
    ("audit_pure.py",
     "GDAL-free self-audit — naming, DGED tables, version consistency across all "
     "12 declarations"),
    ("tests/",
     "pytest suite (conftest, test_lib, test_validator, test_converters). Run "
     "'pytest' from the project root"),
    ("RELEASE_CHECK_v0.46.py",
     "Release gate — audit, pytest, real conversions, validation, ASCII-console "
     "check, PyInstaller build and run"),
    ("PACKAGE_v0.46.py",
     "Builds the release zips (full tool + validator-only bundle)"),
    ("START_HERE.md",
     "One-page orientation — read this first"),
    ("VERSION.txt",
     "Full changelog in prose. Hand-maintained below the header: the packagers "
     "rewrite only the three header lines"),
]

# (heading, symptom segments, [bullet segments, ...])
TROUBLESHOOTING = [
    (
        "Wrong Python interpreter (ModuleNotFoundError inside an activated environment)",
        [("ModuleNotFoundError: No module named 'numpy'", "code"),
         (" or ", "body"),
         ("No module named 'osgeo'", "code"),
         (" — while the prompt clearly shows the environment is active.", "body")],
        [
            [("Cause: the command form, not the environment. Typed as ", "body"),
             ("script.py", "code"),
             (" rather than ", "body"),
             ("python script.py", "code"),
             (", Windows follows the .py file association and runs a different "
              "interpreter. conda activate changes PATH; it does not change the "
              "file association, so the prompt still shows the environment name "
              "while the script runs elsewhere.", "body")],
            [("Always type python first: ", "body"),
             ("python audit_pure.py", "code"),
             (", not ", "body"),
             ("audit_pure.py", "code"),
             (".", "body")],
            [("To confirm which interpreter you are on: ", "body"),
             ("python dem2dged_env.py", "code"),
             (" — it prints sys.executable, CONDA_PREFIX and whether the two "
              "agree.", "body")],
            [("Do not reinstall the packages. They were never missing; this was "
              "misdiagnosed twice before v0.46 added the diagnostic.", "body")],
        ],
    ),
    (
        "UnicodeEncodeError on a non-UTF-8 console (cp949, cp932, cp936)",
        [("UnicodeEncodeError: 'cp949' codec can't encode character", "code"),
         (" during packaging or verification.", "body")],
        [
            [("Fixed in v0.44: every console-facing script now prints pure ASCII "
              "([OK] / [FAIL] / [WARN]) instead of decorative check-mark glyphs, "
              "which encode under UTF-8 and not at all under a national code "
              "page.", "body")],
            [("What can still trigger it is an interpolated value — a Korean or "
              "accented directory name in a path. ", "body"),
             ("dem2dged_lib.safe_print()", "code"),
             (" re-encodes through the console's own codec and degrades to ASCII "
              "rather than raising.", "body")],
            [("On an older copy, force a UTF-8 console first: ", "body"),
             ("chcp 65001", "code"),
             (", or set ", "body"),
             ("PYTHONIOENCODING=utf-8", "code"),
             (".", "body")],
        ],
    ),
    (
        "Source CRS has no EPSG code",
        [("ERROR: cannot determine the EPSG code of <file>", "code"),
         (" — the source raster's CRS carries no EPSG authority code (a bare ESRI "
          "WKT, a local or engineering CRS, a plain .asc grid, or no projection "
          "at all).", "body")],
        [
            [("Tag the source with a real EPSG code, then re-run: ", "body"),
             ("python -m osgeo_utils.gdal_edit -a_srs EPSG:4326 my_dem.tif", "code")],
            [("Use the module form. On a standard conda install for Windows the "
              "GDAL utilities ship as modules, not console scripts, so ", "body"),
             ("gdal_edit.py", "code"),
             (" is often not on PATH.", "body")],
            [("Before v0.42 this failed as a TypeError on int(None), naming "
              "neither the file nor the problem; the check now runs while the "
              "filename is still in scope.", "body")],
        ],
    ),
]

VERSION_ROWS = [
    ("v0.46", "2026-08-10",
     "Packaging no longer overwrites VERSION.txt's maintained changelog with a "
     "stale hardcoded copy (which is why the v0.41 notes exist nowhere). New "
     "dem2dged_env.py diagnoses the wrong-interpreter case that the tool's own "
     "error messages had been misdiagnosing as a broken environment."),
    ("v0.44", "2026-08-10",
     "Fixed a release blocker on a Korean Windows console (cp949): decorative "
     "glyphs in console output raised UnicodeEncodeError, and the except handler "
     "printed a glyph too, destroying the original error. All console output is "
     "ASCII; new safe_print(); new release-gate step runs every script under "
     "PYTHONIOENCODING=ascii:strict."),
    ("v0.43", "2026-08-10",
     "Closed a real coverage hole: the north/south row pass of tile-edge "
     "reconciliation had never executed in any test, and the test that appeared "
     "to cover it always took a pytest.skip branch. New 2x2 level-0 fixture "
     "covers both passes, Int16 output and the level 0-3 filename form; the "
     "release gate now builds and runs the frozen exes."),
    ("v0.42", "2026-08-09",
     "Release-gate pass over v0.41. Blocker: the packager's EXCLUDE_DIRS listed "
     "\"tests\", so packaging stripped the test suite out of every release zip. "
     "Pre-flight guards added for an unknown -resample value, a missing "
     "gdalwarp, and a source CRS with no EPSG code; a run that produces no tiles "
     "is now a hard error instead of a reported success."),
    ("v0.41", "2026-08-07",
     "Repair release: the v0.40 cut of dem2dged_validate.py was missing an entire "
     "block of code and did not compile, so validation silently never ran, the "
     "GUI checkbox was disabled and the self-audit could not start. Restored, "
     "along with the test suite and a version audit that had been matching "
     "nothing. No change to tile geometry, filenames or metadata."),
]


# -- main ---------------------------------------------------------------------

def main():
    if not os.path.exists(DOCX):
        sys.stderr.write("ERROR: not found: %s\n" % DOCX)
        return 2

    # The ".bak" suffix is deliberate: dem2dged_package.py's
    # EXCLUDE_FILE_SUFFIXES already skips it, so the backup cannot end up
    # inside a release zip the way the stray lu49gpd00.tmp did before v0.42.
    backup = DOCX.replace(".docx", "_v0.40_backup.docx.bak")
    if not os.path.exists(backup):
        shutil.copy2(DOCX, backup)
        print("[OK]   backup written: %s" % os.path.basename(backup))

    doc = Document(DOCX)
    changed = 0

    # 1. cover -----------------------------------------------------------------
    cover = doc.tables[T_COVER_VERSION].rows[0].cells
    if "0.46" not in cover[0].text:
        set_cell_text(cover[0], "VERSION %s" % NEW_VERSION)
        set_cell_text(cover[1], "%s   ·   Beta release" % NEW_DATE)
        changed += 1
        print("[OK]   cover      -> VERSION %s, %s" % (NEW_VERSION, NEW_DATE))
    else:
        print("[SKIP] cover already at 0.46")

    # 2. footer ----------------------------------------------------------------
    hit = 0
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer,
                       section.even_page_footer):
            if footer is None:
                continue
            for para in footer.paragraphs:
                for run in para.runs:
                    if "0.40-beta" in run.text:
                        run.text = run.text.replace("0.40-beta", NEW_VERSION)
                        hit += 1
    if hit:
        changed += 1
        print("[OK]   footer     -> v%s (%d run(s))" % (NEW_VERSION, hit))
    else:
        print("[SKIP] footer already current")

    # 3. package contents ------------------------------------------------------
    table = doc.tables[T_PACKAGE_CONTENTS]
    existing = {row.cells[0].text.strip() for row in table.rows}
    added = 0
    for name, purpose in PACKAGE_ROWS:
        if name in existing:
            continue
        append_row(table, 1, [name, purpose])
        added += 1
    if added:
        changed += 1
    print("[%s] package contents: %d row(s) added"
          % ("OK" if added else "SKIP", added))

    # 4. troubleshooting -------------------------------------------------------
    anchor = para_index_by_text(doc, "GDAL / PROJ initialization errors")
    if anchor < 0:
        sys.stderr.write("ERROR: troubleshooting anchor heading not found\n")
        return 3
    head_tpl = doc.paragraphs[anchor]        # Heading 3
    symp_tpl = doc.paragraphs[anchor + 1]    # "Symptom:  ..." line
    bull_tpl = doc.paragraphs[anchor + 2]    # bulleted remedy

    existing_headings = {p.text.strip() for p in doc.paragraphs}
    added = 0
    for heading, symptom, bullets in TROUBLESHOOTING:
        if heading in existing_headings:
            continue
        block = [build_paragraph(head_tpl, [(heading, "body")], 0, 0)]
        block.append(build_paragraph(symp_tpl, symptom, 2, 1,
                                     lead_run_idx=0, lead_text="Symptom:  "))
        for bullet in bullets:
            block.append(build_paragraph(bull_tpl, bullet, 0, 1))
        ref = head_tpl._p
        for element in block:
            ref.addprevious(element)
        added += 1
    if added:
        changed += 1
    print("[%s] troubleshooting: %d entr(y/ies) added"
          % ("OK" if added else "SKIP", added))

    # 5. version history -------------------------------------------------------
    table = doc.tables[T_VERSION_HISTORY]
    existing = {row.cells[0].text.strip() for row in table.rows}
    added = 0
    for version, date, text in reversed(VERSION_ROWS):
        if version in existing:
            continue
        clone_row_after(table, 1, 0, [version, date, text])
        added += 1
    if added:
        changed += 1
    print("[%s] version history: %d row(s) added"
          % ("OK" if added else "SKIP", added))

    if not changed:
        print("\n[OK] nothing to do -- the manual is already at v%s" % NEW_VERSION)
        return 0

    doc.save(DOCX)
    print("\n[OK] saved %s at %s"
          % (os.path.basename(DOCX), datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
    print("     NOTE: the contents list on page 2 is static text. If the page\n"
          "     numbers for sections 10 and 11 shifted, correct them by hand.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
