#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# SPDX-License-Identifier: GPL-2.0-or-later
# Copyright (c) 2026 Eui Soo SON
"""
update_manual_v0.49.py -- bring DEM2DGED_User_Manual.docx up to v0.49.

Version: 0.49

WHY THIS EXISTS
---------------
Same purpose as update_manual_v0.47.py (and v0.46 before it): audit_pure.py
enforces version consistency across the twelve *code* declarations, but
neither it nor RELEASE_CHECK looks at README.md, QUICKSTART.html or this
manual, so the manual can silently fall behind again with nothing to flag
it.

v0.48 itself makes no change to tile geometry, resampling algorithms, or
metadata -- README.md and QUICKSTART.html were updated to explain why
elevation accuracy degrades faster on high-relief (mountainous) terrain
than on flat terrain at a fixed DGED level, to show the DGED spec's
predicted horizontal/vertical accuracy goal per level next to GSD in the
product-levels table, and to recommend "-resample optimize" for
high-relief sources. This script carries the same explanation into the
manual as a new FAQ-style entry, using the exact same paragraph-cloning
mechanism the troubleshooting section already relies on
(build_paragraph() / para_index_by_text()) rather than guessing at a new
table's layout.

As a same-pass follow-up, two of the project's OTHER versioned filenames
were also renamed to catch them up: RELEASE_CHECK_v0.47.py ->
RELEASE_CHECK_v0.48.py and PACKAGE_v0.47.py -> PACKAGE_v0.48.py (internal
text and EXPECTED_VERSION updated to match), plus the
dem2dged_validate_v0.47/ validator-only bundle folder ->
dem2dged_validate_v0.48/. See step 3 below for what that means for the
manual's package-contents table specifically.

NOT EXECUTED OR VERIFIED THIS SESSION -- READ THIS FIRST
----------------------------------------------------------
Exactly the same caveat as update_manual_v0.47.py: this script was written
and syntax-checked in an environment that does not have
DEM2DGED_User_Manual.docx available (it is not part of the mounted project
folder), so it has NOT been run against the real file. It is modelled
closely on update_manual_v0.47.py, which in turn was modelled on the v0.46
script that WAS run and verified -- but treat this one as a draft to review
before trusting it, and read the printed [OK]/[SKIP] lines when you do run
it.

Because the actual current state of your manual is unknown from here (it
may be at "0.47-beta" if update_manual_v0.47.py was run, or an older
predecessor if it was not), this script checks EACH predecessor string in
turn and upgrades whichever it finds -- see "1. cover" / "2. footer" below.
Every step is still guarded and idempotent: re-running it is a no-op, not a
duplicate.

WHAT IT CHANGES
---------------
  1. Cover block          "VERSION 0.40-beta" / "0.46-beta" / "0.47-beta" / "0.48-beta"
                          -> "VERSION 0.49-beta"; date -> "August 2026"
  2. Page footer          same predecessor set -> "0.49-beta"
  3. Package contents     PACKAGE_ROWS now names RELEASE_CHECK_v0.48.py /
                          PACKAGE_v0.48.py (the current, real filenames)
                          instead of the v0.47 names. CAVEAT: the dedup
                          check in step 3 below matches by EXACT row name,
                          so if update_manual_v0.47.py already ran against
                          this file and added rows named
                          "RELEASE_CHECK_v0.47.py" / "PACKAGE_v0.47.py",
                          this script does NOT remove or rename those old
                          rows -- it only adds new ones alongside them,
                          since python-docx table editing here is
                          append-only by design (see clone_row_after()).
                          If your manual already has the v0.47-named rows,
                          delete those two rows by hand after running this
                          script.

  4. Troubleshooting/FAQ  unchanged 4 entries from v0.46/v0.47/v0.48, PLUS
                          one new v0.49 FAQ entry: "Converted mountain tiles
                          show ridges or hollows that are not in the source"
                          -- explains aliasing, why no resampler can fix it
                          afterwards, the new --prefilter gaussian flag, and
                          the measured terrain classes where it HURTS.
                          The v0.48 entry "Elevation errors vary a lot
                          between flat and mountainous areas" -- explains
                          post spacing vs. resampling, the measured
                          Bilinear/Cubic vs. Nearest Neighbor gap on steep
                          terrain, and the two levers (matching level,
                          "-resample optimize") that actually help.
  5. Version history      adds a row for v0.49 (cumulative back through
                          v0.41, same reasoning as v0.47's script)

USAGE (Anaconda Prompt, dedicated environment -- never base):
    conda activate dem2dged_anaconda_environment
    pip install python-docx
    python update_manual_v0.49.py

Note the command form: "python update_manual_v0.49.py", not
"update_manual_v0.49.py". See dem2dged_env.py for why that matters.
"""

import copy
import os
import shutil
import sys
from datetime import datetime

try:
    from docx import Document
except ImportError:
    sys.stderr.write(
        "ERROR: python-docx is not available in the interpreter that is running\n"
        "       this script (%s).\n"
        "       Install it with:  pip install python-docx\n"
        "       If you believe it IS installed, run 'python dem2dged_env.py'\n"
        "       first -- you are probably on a different interpreter than you\n"
        "       think.\n" % sys.executable
    )
    raise SystemExit(2)

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(HERE, "DEM2DGED_User_Manual.docx")
NEW_VERSION = "0.49-beta"
NEW_DATE = "August 2026"
PREDECESSOR_VERSIONS = ("0.48-beta", "0.47-beta", "0.46-beta", "0.40-beta")   # checked in this order

# Table indices, confirmed by inspection of the v0.40 manual (unchanged
# through v0.47 -- every prior script only appends rows/paragraphs, none
# ever restructures a table, so these should still hold).
T_COVER_VERSION = 2
T_PACKAGE_CONTENTS = 5
T_VERSION_HISTORY = 30


# -- helpers ------------------------------------------------------------------

def set_cell_text(cell, text):
    """Replace a cell's text, keeping the formatting of its first run."""
    para = cell.paragraphs[0]
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for run in para.runs[1:]:
        run._r.getparent().remove(run._r)
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)


def clone_row_after(table, template_index, insert_after_index, values):
    """Deep-copy a row for its formatting, fill it, and place it in the table."""
    new_tr = copy.deepcopy(table.rows[template_index]._tr)
    table.rows[insert_after_index]._tr.addnext(new_tr)
    row = table.rows[insert_after_index + 1]
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value)
    return row


def append_row(table, template_index, values):
    return clone_row_after(table, template_index, len(table.rows) - 1, values)


def build_paragraph(template_p, segments, body_run_idx, code_run_idx,
                    lead_run_idx=None, lead_text=None):
    """Clone a paragraph and rebuild its runs from (text, kind) segments.

    kind is "body" for prose or "code" for a monospaced inline snippet; the
    run formatting for each is taken from the template paragraph itself, so
    the result matches the surrounding document exactly.
    """
    new_p = copy.deepcopy(template_p._p)
    runs = list(new_p.findall(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"))
    body_tpl = copy.deepcopy(runs[body_run_idx])
    code_tpl = copy.deepcopy(runs[code_run_idx])
    lead_tpl = copy.deepcopy(runs[lead_run_idx]) if lead_run_idx is not None else None
    for run in runs:
        new_p.remove(run)

    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    def emit(template, text):
        run = copy.deepcopy(template)
        t = run.find(W + "t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        new_p.append(run)

    if lead_tpl is not None and lead_text is not None:
        emit(lead_tpl, lead_text)
    for text, kind in segments:
        emit(code_tpl if kind == "code" else body_tpl, text)
    return new_p


def para_index_by_text(doc, needle):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == needle:
            return i
    return -1


# -- content ------------------------------------------------------------------

# Cumulative: every row added since v0.40, unchanged from update_manual_v0.47.py.
# v0.49 adds one new package file (selftest_prefilter.py) and renames the
# release-gate/packaging scripts to the v0.49 filenames. Kept cumulative so
# this script is still safe to run against a manual at any older state.
PACKAGE_ROWS = [
    ("dem2dged_compare.py",
     "Resampling comparison test — ranks Nearest / Bilinear / Cubic against the "
     "source and writes an HTML report"),
    ("dem2dged_env.py",
     "Environment diagnostic (v0.46) — dependency-free; run it when an import "
     "fails to see which interpreter is actually running"),
    ("audit_pure.py",
     "GDAL-free self-audit — naming, DGED tables, version consistency across all "
     "12 declarations"),
    ("tests/",
     "pytest suite (conftest, test_lib, test_validator, test_converters). Run "
     "'pytest' from the project root"),
    ("RELEASE_CHECK_v0.49.py",
     "Release gate — audit, pytest, real conversions, validation, ASCII-console "
     "check, PyInstaller build and run"),
    ("PACKAGE_v0.49.py",
     "Builds the release zips (full tool + validator-only bundle)"),
    ("START_HERE.md",
     "One-page orientation — read this first"),
    ("VERSION.txt",
     "Full changelog in prose. Hand-maintained below the header: the packagers "
     "rewrite only the three header lines"),
    ("DIAGNOSE_SECTION_H_v0.11.py",
     "Standalone, read-only diagnostic for a Section H/H2 FAIL — measures the "
     "source and delivered tiles several ways and explains whether a min/max/"
     "mean gap is a measurement artifact or a real defect. Run with --selftest "
     "for a no-data-required self-check"),
    ("selftest_optimize_resampling.py",
     "End-to-end self-test of -resample optimize's clamp-fairness fix on a "
     "synthetic cliff DEM"),
    ("selftest_prefilter.py",
     "Self-test of the v0.49 anti-alias pre-filter -- measures the error "
     "reduction across six terrain-roughness classes on fractal surfaces, "
     "including the near-planar case where the filter makes things worse"),
]

# (heading, symptom segments, [bullet segments, ...]).
# The first 3 entries are unchanged from v0.46/v0.47 -- kept here (rather than
# imported) so this script has no dependency on either predecessor and stays
# idempotent against a manual at any of the older baseline states. The 4th
# entry is new in v0.48.
TROUBLESHOOTING = [
    (
        "Wrong Python interpreter (ModuleNotFoundError inside an activated environment)",
        [("ModuleNotFoundError: No module named 'numpy'", "code"),
         (" or ", "body"),
         ("No module named 'osgeo'", "code"),
         (" — while the prompt clearly shows the environment is active.", "body")],
        [
            [("Cause: the command form, not the environment. Typed as ", "body"),
             ("script.py", "code"),
             (" rather than ", "body"),
             ("python script.py", "code"),
             (", Windows follows the .py file association and runs a different "
              "interpreter. conda activate changes PATH; it does not change the "
              "file association, so the prompt still shows the environment name "
              "while the script runs elsewhere.", "body")],
            [("Always type python first: ", "body"),
             ("python audit_pure.py", "code"),
             (", not ", "body"),
             ("audit_pure.py", "code"),
             (".", "body")],
            [("To confirm which interpreter you are on: ", "body"),
             ("python dem2dged_env.py", "code"),
             (" — it prints sys.executable, CONDA_PREFIX and whether the two "
              "agree.", "body")],
            [("Do not reinstall the packages. They were never missing; this was "
              "misdiagnosed twice before v0.46 added the diagnostic.", "body")],
        ],
    ),
    (
        "UnicodeEncodeError on a non-UTF-8 console (cp949, cp932, cp936)",
        [("UnicodeEncodeError: 'cp949' codec can't encode character", "code"),
         (" during packaging or verification.", "body")],
        [
            [("Fixed in v0.44: every console-facing script now prints pure ASCII "
              "([OK] / [FAIL] / [WARN]) instead of decorative check-mark glyphs, "
              "which encode under UTF-8 and not at all under a national code "
              "page.", "body")],
            [("What can still trigger it is an interpolated value — a Korean or "
              "accented directory name in a path. ", "body"),
             ("dem2dged_lib.safe_print()", "code"),
             (" re-encodes through the console's own codec and degrades to ASCII "
              "rather than raising.", "body")],
            [("On an older copy, force a UTF-8 console first: ", "body"),
             ("chcp 65001", "code"),
             (", or set ", "body"),
             ("PYTHONIOENCODING=utf-8", "code"),
             (".", "body")],
        ],
    ),
    (
        "Source CRS has no EPSG code",
        [("ERROR: cannot determine the EPSG code of <file>", "code"),
         (" — the source raster's CRS carries no EPSG authority code (a bare ESRI "
          "WKT, a local or engineering CRS, a plain .asc grid, or no projection "
          "at all).", "body")],
        [
            [("Tag the source with a real EPSG code, then re-run: ", "body"),
             ("python -m osgeo_utils.gdal_edit -a_srs EPSG:4326 my_dem.tif", "code")],
            [("Use the module form. On a standard conda install for Windows the "
              "GDAL utilities ship as modules, not console scripts, so ", "body"),
             ("gdal_edit.py", "code"),
             (" is often not on PATH.", "body")],
            [("Before v0.42 this failed as a TypeError on int(None), naming "
              "neither the file nor the problem; the check now runs while the "
              "filename is still in scope.", "body")],
        ],
    ),
    # -- new in v0.48 --
    (
        "Elevation errors vary a lot between flat and mountainous areas",
        [("A delivery converted from flat terrain compares to the source with a "
          "small residual; the same level and settings on mountainous terrain "
          "show a much larger sample-window difference.", "body")],
        [
            [("This is expected, not a bug. A level's GSD is the spacing "
              "", "body"),
             ("between", "code"),
             (" measured posts — every value in between is filled in by "
              "resampling, not measured. On flat terrain neighbouring posts are "
              "nearly equal, so any resampler reconstructs the surface between "
              "them closely. On steep terrain (ridgelines, canyon walls, cliff "
              "edges) elevation can change tens of metres within a single post "
              "spacing, so the resampler is effectively inventing values rather "
              "than interpolating them.", "body")],
            [("Measured on real steep terrain: Bilinear and Cubic Convolution "
              "showed 6–10 m sample-window error, while Nearest Neighbor — which "
              "copies the nearest post instead of blending across the slope — "
              "stayed near 0 m on the same source (see the v0.46 changelog "
              "entry in VERSION.txt). Nearest Neighbor is not more accurate in "
              "general; it is the only one of the three that can represent a "
              "true discontinuity without smoothing across it.", "body")],
            [("Fix 1 — pick the level that matches your source data's real "
              "resolution, not a higher one. Level only controls output post "
              "spacing; it cannot add detail your source DEM never had.", "body")],
            [("Fix 2 — use ", "body"),
             ("-resample optimize", "code"),
             (" (GUI: \"Optimize\") on high-relief sources. Since v0.36 it "
              "measures Nearest Neighbor, Bilinear, Cubic Convolution and Cubic "
              "B-Spline directly against your source DEM and uses whichever "
              "reconstructs it most accurately for that specific file, instead "
              "of a fixed rule of thumb.", "body")],
            [("To see the comparison yourself: the GUI's \"Resampling "
              "Comparison Test\" runs any subset of the three manual methods "
              "side by side and writes a ranked HTML report.", "body")],
        ],
    ),
    # -- new in v0.49 --
    (
        "Converted mountain tiles show ridges or hollows that are not in the source",
        [("A delivery downsampled from a fine source to a coarser level contains "
          "plausible-looking terrain structure -- ridges, hollows, banding -- in "
          "places where the source DEM has nothing of the sort. Changing the "
          "resampling method moves it around but never removes it.", "body")],
        [
            [("This is ALIASING, and it is a different problem from resampling "
              "error. Downsampling is decimation, and terrain detail with a "
              "wavelength shorter than twice the target post spacing cannot be "
              "represented at that spacing by any resampler that will ever "
              "exist. It does not disappear quietly -- it folds back into the "
              "product at a longer wavelength the grid CAN represent, which is "
              "why it looks like terrain and why no later choice of resampler "
              "removes it.", "body")],
            [("Mountainous sources are hit hardest because real topography "
              "follows roughly a 1/f^beta power spectrum: rough terrain carries "
              "far more short-wavelength energy (ridge crests, gully walls, "
              "scree, cliff edges) than flat terrain, so there is much more of "
              "it to fold back.", "body")],
            [("Fix (new in v0.49) -- low-pass filter the source BEFORE "
              "decimating, the standard practice everywhere else in signal and "
              "image processing: ", "body"),
             ("--prefilter gaussian", "code"),
             (". The sigma defaults to (target GSD / source GSD - 1) / 2 source "
              "pixels and the filter is skipped entirely when you are not "
              "downsampling. Tune it with ", "body"),
             ("--prefilter-sigma", "code"),
             (".", "body")],
            [("It is OFF by default, on purpose. Measured on fractal terrain "
              "(2 m -> 20 m posts, scored against an ideal band-limited "
              "reference): very rough -82.0%, mountainous -73.8%, hilly -63.4%, "
              "rolling -51.0% error -- but smooth terrain only -16.4%, and "
              "NEAR-PLANAR TERRAIN GETS 72.9% WORSE. Where there is no "
              "short-wavelength energy there is nothing to alias, so the filter "
              "can only blur real signal.", "body")],
            [("It is also a bias/variance trade: low-pass filtering "
              "systematically clips real summits and fills real valleys. If your "
              "acceptance criterion is LE90 against surveyed spot heights on "
              "peaks, weigh that against the aliasing before using it.", "body")],
            [("So measure it on your own data: convert the same source twice, "
              "once with and once without the flag, and compare the two "
              "DGED_Validation_Report.html files. Run ", "body"),
             ("python selftest_prefilter.py", "code"),
             (" to reproduce the table above on your own GDAL build.", "body")],
        ],
    ),
]

# Cumulative for the same reason as PACKAGE_ROWS above.
VERSION_ROWS = [
    ("v0.49", "2026-08-13",
     "New opt-in Gaussian anti-alias pre-filter for high-relief sources "
     "(--prefilter gaussian / --prefilter-sigma). Default is none, which "
     "reproduces v0.48 output bit for bit -- an existing delivery does not "
     "need regenerating. Downsampling a DEM is decimation, and terrain detail "
     "shorter than twice the target post spacing cannot be represented at "
     "that spacing by any resampler; without a low-pass filter first it folds "
     "back into the product as false longer-wavelength structure (aliasing) "
     "that looks like terrain but is not. Rough terrain carries far more "
     "short-wavelength energy than flat terrain, so mountainous sources are "
     "hit hardest -- a distinct error mechanism from the interpolation error "
     "v0.48 documented. Measured on fractal terrain, 2 m -> 20 m posts, "
     "scored against an ideal band-limited reference: very rough -82.0%, "
     "mountainous -73.8%, hilly -63.4%, rolling -51.0%, smooth -16.4%, and "
     "near-planar 72.9% WORSE -- which is why the feature is opt-in rather "
     "than automatic. NoData is handled by normalised convolution (a plain "
     "convolution smears the -32767 sentinel into real terrain), void "
     "footprints are preserved exactly, the sigma used is recorded in the "
     "sidecar lineage, and tile geometry, grid snapping, resampler choice and "
     "the cubic-family clamp are all untouched. New selftest_prefilter.py."),
    ("v0.48", "2026-08-12",
     "No change to tile geometry, resampling algorithms or metadata. "
     "README.md, QUICKSTART.html and this manual now explain why elevation "
     "accuracy degrades faster on high-relief (mountainous) terrain than "
     "on flat terrain at a fixed DGED level -- a coarse post spacing "
     "under-samples short-wavelength relief, and every interpolation-based "
     "resampler assumes the terrain is locally smooth between posts, an "
     "assumption steep slopes violate. The product-levels table now lists "
     "the DGED spec's predicted horizontal (CE90) / vertical (LE90) "
     "accuracy goal per level next to GSD. Both documents recommend "
     "-resample optimize for high-relief sources. Also fixed a "
     "QUICKSTART.html gap: its options table listed -resample's valid "
     "values without \"optimize\", though it has been valid since v0.36. "
     "Follow-up in the same pass: RELEASE_CHECK_v0.47.py / PACKAGE_v0.47.py "
     "renamed to the v0.48 filenames (EXPECTED_VERSION and internal text "
     "updated to match), dem2dged_validate_v0.47/ renamed to "
     "dem2dged_validate_v0.48/ with its internal file copies refreshed, "
     "and version_info_gui.txt / version_info_validate.txt regenerated "
     "from the stale 0.45 they were stuck at."),
    ("v0.47", "2026-08-12",
     "Fixed a fairness bug in -resample optimize's accuracy measurement: "
     "cubic-family candidates were scored on their raw, unclamped "
     "reconstruction, while delivered tiles made with those algorithms are "
     "clamped into the source's true range before being written. A few "
     "overshoot pixels at sharp discontinuities -- never actually shipped -- "
     "could inflate RMSE and make cubic-family methods look worse than what "
     "a user actually receives. _holdout_stats() now applies the identical "
     "clamp before scoring. Cubic B-Spline added as a fourth measured "
     "candidate, safe now that overshoot is clamped. Also fixes a "
     "version-consistency gap: dem2dged_package.py, "
     "dem2dged_validate_package.py and BUILD_AND_PACKAGE.py had been stuck "
     "at VERSION = \"0.45\" since the v0.46 release."),
    ("v0.46", "2026-08-10",
     "Packaging no longer overwrites VERSION.txt's maintained changelog with a "
     "stale hardcoded copy (which is why the v0.41 notes exist nowhere). New "
     "dem2dged_env.py diagnoses the wrong-interpreter case that the tool's own "
     "error messages had been misdiagnosing as a broken environment."),
    ("v0.44", "2026-08-10",
     "Fixed a release blocker on a Korean Windows console (cp949): decorative "
     "glyphs in console output raised UnicodeEncodeError, and the except handler "
     "printed a glyph too, destroying the original error. All console output is "
     "ASCII; new safe_print(); new release-gate step runs every script under "
     "PYTHONIOENCODING=ascii:strict."),
    ("v0.43", "2026-08-10",
     "Closed a real coverage hole: the north/south row pass of tile-edge "
     "reconciliation had never executed in any test, and the test that appeared "
     "to cover it always took a pytest.skip branch. New 2x2 level-0 fixture "
     "covers both passes, Int16 output and the level 0-3 filename form; the "
     "release gate now builds and runs the frozen exes."),
    ("v0.42", "2026-08-09",
     "Release-gate pass over v0.41. Blocker: the packager's EXCLUDE_DIRS listed "
     "\"tests\", so packaging stripped the test suite out of every release zip. "
     "Pre-flight guards added for an unknown -resample value, a missing "
     "gdalwarp, and a source CRS with no EPSG code; a run that produces no tiles "
     "is now a hard error instead of a reported success."),
    ("v0.41", "2026-08-07",
     "Repair release: the v0.40 cut of dem2dged_validate.py was missing an entire "
     "block of code and did not compile, so validation silently never ran, the "
     "GUI checkbox was disabled and the self-audit could not start. Restored, "
     "along with the test suite and a version audit that had been matching "
     "nothing. No change to tile geometry, filenames or metadata."),
]


# -- main ---------------------------------------------------------------------

def main():
    if not os.path.exists(DOCX):
        sys.stderr.write("ERROR: not found: %s\n" % DOCX)
        return 2

    # The ".bak" suffix is deliberate: dem2dged_package.py's
    # EXCLUDE_FILE_SUFFIXES already skips it, so the backup cannot end up
    # inside a release zip the way the stray lu49gpd00.tmp did before v0.42.
    # Kept as the ORIGINAL "_v0.40_backup" name on purpose -- one baseline
    # reference regardless of which version script runs, not a new backup
    # per release.
    backup = DOCX.replace(".docx", "_v0.40_backup.docx.bak")
    if not os.path.exists(backup):
        shutil.copy2(DOCX, backup)
        print("[OK]   backup written: %s" % os.path.basename(backup))

    doc = Document(DOCX)
    changed = 0

    # 1. cover -----------------------------------------------------------------
    cover = doc.tables[T_COVER_VERSION].rows[0].cells
    if "0.49" not in cover[0].text:
        set_cell_text(cover[0], "VERSION %s" % NEW_VERSION)
        set_cell_text(cover[1], "%s   ·   Beta release" % NEW_DATE)
        changed += 1
        print("[OK]   cover      -> VERSION %s, %s" % (NEW_VERSION, NEW_DATE))
    else:
        print("[SKIP] cover already at %s" % NEW_VERSION)

    # 2. footer ------------------------------------------------------------
    # Unknown from here which predecessor script(s) already ran against this
    # file, so every possible predecessor string is checked.
    hit = 0
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer,
                       section.even_page_footer):
            if footer is None:
                continue
            for para in footer.paragraphs:
                for run in para.runs:
                    for predecessor in PREDECESSOR_VERSIONS:
                        if predecessor in run.text:
                            run.text = run.text.replace(predecessor, NEW_VERSION)
                            hit += 1
                            break
    if hit:
        changed += 1
        print("[OK]   footer     -> v%s (%d run(s))" % (NEW_VERSION, hit))
    else:
        print("[SKIP] footer already current (or predecessor string not found "
              "-- check by hand if this is unexpected)")

    # 3. package contents ------------------------------------------------------
    table = doc.tables[T_PACKAGE_CONTENTS]
    existing = {row.cells[0].text.strip() for row in table.rows}
    added = 0
    for name, purpose in PACKAGE_ROWS:
        if name in existing:
            continue
        append_row(table, 1, [name, purpose])
        added += 1
    if added:
        changed += 1
    print("[%s] package contents: %d row(s) added"
          % ("OK" if added else "SKIP", added))

    # 4. troubleshooting / FAQ ---------------------------------------------------
    anchor = para_index_by_text(doc, "GDAL / PROJ initialization errors")
    if anchor < 0:
        sys.stderr.write("ERROR: troubleshooting anchor heading not found\n")
        return 3
    head_tpl = doc.paragraphs[anchor]        # Heading 3
    symp_tpl = doc.paragraphs[anchor + 1]    # "Symptom:  ..." line
    bull_tpl = doc.paragraphs[anchor + 2]    # bulleted remedy

    existing_headings = {p.text.strip() for p in doc.paragraphs}
    added = 0
    for heading, symptom, bullets in TROUBLESHOOTING:
        if heading in existing_headings:
            continue
        block = [build_paragraph(head_tpl, [(heading, "body")], 0, 0)]
        block.append(build_paragraph(symp_tpl, symptom, 2, 1,
                                     lead_run_idx=0, lead_text="Symptom:  "))
        for bullet in bullets:
            block.append(build_paragraph(bull_tpl, bullet, 0, 1))
        ref = head_tpl._p
        for element in block:
            ref.addprevious(element)
        added += 1
    if added:
        changed += 1
    print("[%s] troubleshooting/FAQ: %d entr(y/ies) added"
          % ("OK" if added else "SKIP", added))

    # 5. version history -------------------------------------------------------
    table = doc.tables[T_VERSION_HISTORY]
    existing = {row.cells[0].text.strip() for row in table.rows}
    added = 0
    for version, date, text in reversed(VERSION_ROWS):
        if version in existing:
            continue
        clone_row_after(table, 1, 0, [version, date, text])
        added += 1
    if added:
        changed += 1
    print("[%s] version history: %d row(s) added"
          % ("OK" if added else "SKIP", added))

    if not changed:
        print("\n[OK] nothing to do -- the manual is already at v%s" % NEW_VERSION)
        return 0

    doc.save(DOCX)
    print("\n[OK] saved %s at %s"
          % (os.path.basename(DOCX), datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
    print("     NOTE: the contents list on page 2 is static text. If the page\n"
          "     numbers for sections 10 and 11 (or wherever the new FAQ entry\n"
          "     landed) shifted, correct them by hand.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
